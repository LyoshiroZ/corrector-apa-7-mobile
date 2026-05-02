"""Corrector APA 7 - aplicación móvil (Kivy + KivyMD).

Versión móvil de la app Streamlit, lista para empaquetar como APK/AAB con
buildozer. Conserva el prompt, parser, métricas, generador de Word y la
opción de keep_with_next para los títulos.
"""

from __future__ import annotations

import io
import json
import os
import re
import sys
import threading
import traceback
import difflib
from collections import Counter
from pathlib import Path


# --- Capturador de errores de emergencia (para diagnostico en Android) -------
# v13: ahora ademas escribimos un "rastro" del arranque (boot trace) para saber
# cuanto progreso de la app ocurre antes del crash. Tambien agregamos timestamp
# y version para distinguir logs viejos de nuevos.
APP_VERSION = "1.0.3"


def _carpeta_logs() -> str:
    """Devuelve carpeta donde escribir logs (Downloads en Android, home en PC)."""
    try:
        from kivy.utils import platform as _plt
        if _plt == "android":
            from android.storage import primary_external_storage_path
            return os.path.join(primary_external_storage_path(), "Download")
    except Exception:
        pass
    return os.path.expanduser("~")


def _escribir_archivo(nombre: str, contenido: str):
    """Escribe un archivo en la carpeta de logs. Silencioso si falla."""
    try:
        ruta = os.path.join(_carpeta_logs(), nombre)
        os.makedirs(os.path.dirname(ruta), exist_ok=True)
        with open(ruta, "w", encoding="utf-8") as f:
            f.write(contenido)
    except Exception:
        pass


def _ahora() -> str:
    import datetime
    return datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")


def _marca_arranque(etapa: str, extra: str = ""):
    """Escribe un archivo con la etapa actual del arranque."""
    txt = f"APU-APA v{APP_VERSION}\nEtapa: {etapa}\nHora: {_ahora()}\n"
    if extra:
        txt += f"\nDetalle:\n{extra}\n"
    _escribir_archivo(f"apa_boot_{etapa}.log", txt)


# Marca 1: Python arrancó y llegamos al inicio de main.py
_marca_arranque("1_python_iniciado")


def _escribir_crash(exc_type, exc_value, exc_tb):
    txt = (
        f"=== APA Crash Log ===\n"
        f"Version APP: {APP_VERSION}\n"
        f"Hora: {_ahora()}\n"
        f"Error: {exc_type.__name__}: {exc_value}\n\n"
    )
    try:
        buf = io.StringIO()
        traceback.print_exception(exc_type, exc_value, exc_tb, file=buf)
        txt += buf.getvalue()
    except Exception:
        pass
    _escribir_archivo("apa_crash.log", txt)


def _excepthook(exc_type, exc_value, exc_tb):
    _escribir_crash(exc_type, exc_value, exc_tb)
    sys.__excepthook__(exc_type, exc_value, exc_tb)


sys.excepthook = _excepthook

from kivy.clock import Clock, mainthread
from kivy.core.window import Window
from kivy.metrics import dp
from kivy.utils import platform

from kivymd.app import MDApp
from kivymd.uix.boxlayout import MDBoxLayout
from kivymd.uix.button import MDFlatButton, MDIconButton, MDRaisedButton
from kivymd.uix.card import MDCard
from kivymd.uix.dialog import MDDialog
from kivymd.uix.floatlayout import MDFloatLayout
from kivymd.uix.label import MDLabel
from kivymd.uix.menu import MDDropdownMenu
from kivymd.uix.progressbar import MDProgressBar
from kivymd.uix.scrollview import MDScrollView
from kivymd.uix.screenmanager import MDScreenManager
from kivymd.uix.screen import MDScreen
from kivymd.uix.selectioncontrol import MDSwitch
from kivymd.uix.snackbar import Snackbar
from kivymd.uix.tab import MDTabs, MDTabsBase
from kivymd.uix.textfield import MDTextField
from kivymd.uix.toolbar import MDTopAppBar

import requests
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


# Marca 2: Todos los imports terminaron sin errores
_marca_arranque("2_imports_ok")


GEMINI_MODELOS = ["gemini-2.0-flash", "gemini-1.5-flash", "gemini-1.5-pro"]
GEMINI_DEFAULT = "gemini-2.0-flash"
GEMINI_REST = (
    "https://generativelanguage.googleapis.com/v1beta/models/"
    "{modelo}:generateContent?key={key}"
)

PROMPT = """Eres un especialista en normas APA 7ma edición y redacción académica en español.

Tu tarea: corregir el siguiente texto académico ({tipo_doc}) al formato APA 7ma edición.

Debes:
1. Corregir todas las CITAS dentro del texto al formato APA 7 (parentéticas y narrativas, citas directas cortas y largas, citas de paráfrasis).
2. Generar/corregir la lista de REFERENCIAS al final, ordenada alfabéticamente, con sangría francesa implícita y formato APA 7 correcto.
3. Mejorar de forma mínima la redacción solo cuando sea necesario para la coherencia, sin cambiar el contenido ni el significado.
4. Conservar el idioma original (español).
5. Si faltan datos en alguna referencia (ej. año, editorial), márcalo entre corchetes así: [falta año].
6. Estructurar el documento con jerarquías de títulos APA 7. Identifica títulos y subtítulos en el texto (o agrégalos si el documento claramente los necesita) y márcalos al INICIO de su línea con uno de estos prefijos exactos:
   - `[H1] ` para títulos de Nivel 1 (secciones principales: Introducción, Método, Resultados, Discusión, Conclusiones, etc.)
   - `[H2] ` para títulos de Nivel 2 (subsecciones)
   - `[H3] ` para títulos de Nivel 3 (sub-subsecciones)
   No uses estos marcadores en líneas que no sean títulos. No los uses en el cuerpo ni en las referencias.

FORMATO DE SALIDA OBLIGATORIO (en texto plano, sin markdown, sin ```):

TÍTULO: <un título sugerido para el documento>

CUERPO:
<el texto corregido completo, con las citas en formato APA correcto>

REFERENCIAS:
<lista de referencias en formato APA 7, una por línea, ordenadas alfabéticamente>

{notas_section}

Texto a corregir:
---
{texto}
---
"""

_STOPWORDS = {
    "de", "la", "que", "el", "en", "y", "a", "los", "del", "se", "las", "por",
    "un", "para", "con", "no", "una", "su", "al", "lo", "como", "más", "pero",
    "sus", "le", "ya", "o", "este", "sí", "porque", "esta", "entre", "cuando",
    "muy", "sin", "sobre", "también", "me", "hasta", "hay", "donde", "quien",
    "desde", "todo", "nos", "durante", "todos", "uno", "les", "ni", "contra",
    "otros", "ese", "eso", "ante", "ellos", "e", "esto", "mí", "antes",
    "algunos", "qué", "unos", "yo", "otro", "otras", "otra", "él", "tanto",
    "esa", "estos", "mucho", "quienes", "nada", "muchos", "cual", "poco",
    "ella", "estar", "estas", "algunas", "algo", "nosotros", "mi", "mis", "tú",
    "te", "ti", "tu", "tus", "ellas", "es", "son", "fue", "ser", "ha", "han",
    "the", "of", "and", "to", "in", "is", "it", "that", "for", "on", "with",
}


# ---------------------------- Helpers de IA ----------------------------------


def llamar_gemini_rest(api_key: str, texto: str, tipo_doc: str, modelo: str,
                       con_notas: bool, timeout: int = 90) -> str:
    """Llama a Gemini por REST con reintentos y manejo de timeouts."""
    if not api_key:
        raise RuntimeError(
            "Configura tu API key de Gemini en el icono de ajustes."
        )
    notas_section = (
        "NOTAS DE CAMBIOS:\n<lista breve de los principales cambios "
        "realizados (máx. 8 puntos)>"
    ) if con_notas else ""
    prompt = PROMPT.format(
        tipo_doc=tipo_doc, notas_section=notas_section, texto=texto,
    )
    url = GEMINI_REST.format(modelo=modelo, key=api_key)
    payload = {
        "contents": [{"parts": [{"text": prompt}]}],
        "generationConfig": {"temperature": 0.2},
    }
    last_err: Exception | None = None
    for intento in range(3):
        try:
            r = requests.post(url, json=payload, timeout=timeout)
            if r.status_code >= 400:
                raise RuntimeError(
                    f"La API respondió con código {r.status_code}. "
                    "Verifica tu API key y conexión."
                )
            data = r.json()
            return (
                data["candidates"][0]["content"]["parts"][0]["text"] or ""
            ).strip()
        except (requests.Timeout, requests.ConnectionError) as e:
            last_err = e
            continue
    raise RuntimeError(
        "No se pudo conectar con Gemini tras 3 intentos. "
        "Revisa tu conexión móvil e inténtalo de nuevo."
    )


def parsear_secciones(salida: str) -> dict:
    secciones = {"titulo": "", "cuerpo": "", "referencias": "", "notas": ""}
    actual: str | None = None
    buffer: list[str] = []

    def flush():
        if actual and buffer:
            secciones[actual] = "\n".join(buffer).strip()

    for linea in salida.splitlines():
        stripped = linea.strip()
        upper = stripped.upper()
        if upper.startswith("TÍTULO:") or upper.startswith("TITULO:"):
            flush(); buffer = []; actual = "titulo"
            valor = stripped.split(":", 1)[1].strip() if ":" in stripped else ""
            if valor:
                buffer.append(valor)
        elif upper.startswith("CUERPO:"):
            flush(); buffer = []; actual = "cuerpo"
        elif upper.startswith("REFERENCIAS:"):
            flush(); buffer = []; actual = "referencias"
        elif upper.startswith("NOTAS DE CAMBIOS:") or upper.startswith("NOTAS:"):
            flush(); buffer = []; actual = "notas"
        else:
            if actual:
                buffer.append(linea)
    flush()
    return secciones


# ---------------------------- Métricas ---------------------------------------


def contar_citas(texto: str) -> int:
    if not texto:
        return 0
    parenteticas = re.findall(
        r"\([^()]*\b\d{4}[a-z]?\b[^()]*\)", texto,
    )
    narrativas = re.findall(
        r"\b[A-ZÁÉÍÓÚÑ][\wÁÉÍÓÚÑáéíóúñ\-]+(?:\s+(?:y|and|&)\s+"
        r"[A-ZÁÉÍÓÚÑ][\wÁÉÍÓÚÑáéíóúñ\-]+)?\s*\(\d{4}[a-z]?\)",
        texto,
    )
    return len(parenteticas) + len(narrativas)


def contar_referencias(texto: str) -> int:
    if not texto:
        return 0
    return sum(1 for l in texto.splitlines() if l.strip())


def contar_cambios_palabras(original: str, corregido: str) -> dict:
    po = (original or "").split()
    pc = (corregido or "").split()
    matcher = difflib.SequenceMatcher(a=po, b=pc, autojunk=False)
    eliminadas = anadidas = modificadas = 0
    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == "delete":
            eliminadas += i2 - i1
        elif tag == "insert":
            anadidas += j2 - j1
        elif tag == "replace":
            modificadas += max(i2 - i1, j2 - j1)
    return {
        "modificadas": modificadas, "anadidas": anadidas,
        "eliminadas": eliminadas,
        "total": modificadas + anadidas + eliminadas,
        "palabras_original": len(po), "palabras_corregido": len(pc),
    }


def _silabas_es(palabra: str) -> int:
    palabra = re.sub(r"[^a-záéíóúüñ]", "", palabra.lower())
    if not palabra:
        return 0
    vocales = "aeiouáéíóúü"
    s = 0
    prev = False
    for ch in palabra:
        v = ch in vocales
        if v and not prev:
            s += 1
        prev = v
    return max(s, 1)


def analizar_legibilidad(texto: str) -> dict:
    texto = texto or ""
    oraciones = [o.strip() for o in re.split(r"[.!?]+", texto) if o.strip()]
    palabras = re.findall(r"[A-Za-zÁÉÍÓÚÜÑáéíóúüñ]+", texto)
    no = max(len(oraciones), 1)
    npal = max(len(palabras), 1)
    nsil = sum(_silabas_es(p) for p in palabras)
    P = nsil / npal
    F = npal / no
    fh = max(0.0, min(100.0, 206.84 - (60 * P) - (1.02 * F)))
    nivel = (
        "Muy fácil" if fh >= 80 else "Fácil" if fh >= 70 else
        "Bastante fácil" if fh >= 60 else "Normal" if fh >= 50 else
        "Algo difícil" if fh >= 40 else "Difícil" if fh >= 30 else
        "Muy difícil"
    )
    significativas = [p.lower() for p in palabras
                      if len(p) > 3 and p.lower() not in _STOPWORDS]
    repetidas = [(p, c) for p, c in Counter(significativas).most_common(8)
                 if c >= 3]
    diversidad = (len(set(p.lower() for p in palabras)) / npal) * 100
    largas = [o for o in oraciones if len(o.split()) > 30]

    consejos: list[str] = []
    if F > 25:
        consejos.append(
            f"Las oraciones son largas (promedio {F:.0f} palabras). "
            "Divide las más extensas para mejorar la lectura."
        )
    if fh < 50:
        consejos.append(
            "La legibilidad es baja. Usa palabras y oraciones más sencillas."
        )
    if diversidad < 45:
        consejos.append(
            f"Diversidad de vocabulario baja ({diversidad:.0f}%). "
            "Apóyate en sinónimos."
        )
    if repetidas:
        top = ", ".join(f"«{p}» ({c})" for p, c in repetidas[:3])
        consejos.append(f"Palabras muy repetidas: {top}.")
    if largas:
        consejos.append(
            f"{len(largas)} oración(es) supera(n) las 30 palabras."
        )
    if not consejos:
        consejos.append(
            "¡Bien hecho! El texto tiene buena legibilidad y variedad léxica."
        )
    return {
        "fh": fh, "nivel": nivel, "F": F, "P": P,
        "diversidad": diversidad, "n_palabras": npal,
        "n_oraciones": no, "repetidas": repetidas,
        "oraciones_largas": len(largas), "consejos": consejos,
    }


def clasificar_referencia(linea: str) -> str:
    t = linea.strip(); tl = t.lower()
    if not t:
        return "Otro"
    if re.search(r"\btesis\b|\bdisertaci[oó]n\b|tesis doctoral|trabajo de grado", tl):
        return "Tesis"
    if re.search(r"\b(actas|conferencia|congreso|simposio|proceedings)\b", tl):
        return "Conferencia"
    if re.search(r"\b(informe|reporte|report)\b", tl) and "doi" not in tl:
        return "Informe"
    if re.search(r"\bdoi\.org\/|https?:\/\/doi\.org|\bdoi:\s*10\.", tl):
        return "Artículo de revista"
    if re.search(r",\s*\d{1,4}\s*\(\s*\d+\s*\)\s*,\s*\d+", t):
        return "Artículo de revista"
    if re.search(r"\ben\s+[A-ZÁÉÍÓÚÑ][^()]+\(eds?\.|\(ed\.\)|\(coords?\.|\(comps?\.", t):
        return "Capítulo de libro"
    if re.search(r"https?:\/\/", tl):
        return "Recurso web"
    if re.search(r"\b(ediciones|editorial|press|publishers?|paid[oó]s|fce|"
                 r"siglo xxi|mcgraw|pearson|alianza|gedisa|anagrama)\b", tl):
        return "Libro"
    return "Otro"


def estadisticas_referencias(referencias: str) -> dict:
    lineas = [l.strip() for l in (referencias or "").splitlines() if l.strip()]
    por_tipo: dict = {}
    por_anio: dict = {}
    sin_anio = 0
    for l in lineas:
        por_tipo[clasificar_referencia(l)] = por_tipo.get(
            clasificar_referencia(l), 0) + 1
        m = re.search(r"\((\d{4})[a-z]?\)", l)
        anio = int(m.group(1)) if m and 1500 <= int(m.group(1)) <= 2100 else None
        if anio is None:
            sin_anio += 1
        else:
            por_anio[anio] = por_anio.get(anio, 0) + 1
    return {
        "total": len(lineas), "por_tipo": por_tipo,
        "por_anio": por_anio, "sin_anio": sin_anio,
    }


# ---------------------------- Word -------------------------------------------


def _aplicar_fuente(run, *, size_pt=12, bold=False, color=None):
    run.font.name = "Times New Roman"
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = rPr.makeelement(qn("w:rFonts"), {})
        rPr.append(rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rFonts.set(qn(attr), "Times New Roman")
    run.font.size = Pt(size_pt)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def _add_parrafo(doc, texto, *, sangria=True, francesa=False, size_pt=12,
                 bold=False, color=None, alineacion=None):
    p = doc.add_paragraph()
    if alineacion is not None:
        p.alignment = alineacion
    pf = p.paragraph_format
    if francesa:
        pf.left_indent = Cm(1.27); pf.first_line_indent = Cm(-1.27)
    elif sangria:
        pf.first_line_indent = Cm(1.27)
    pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    pf.space_after = Pt(0); pf.space_before = Pt(0)
    run = p.add_run(texto)
    _aplicar_fuente(run, size_pt=size_pt, bold=bold, color=color)
    return p


def construir_docx(secciones: dict) -> bytes:
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    rPr = normal.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = rPr.makeelement(qn("w:rFonts"), {})
        rPr.append(rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rFonts.set(qn(attr), "Times New Roman")

    for section in doc.sections:
        section.top_margin = Cm(2.54); section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54); section.right_margin = Cm(2.54)

    titulo = secciones.get("titulo") or "Documento corregido en APA 7"
    _add_parrafo(doc, titulo, sangria=False, bold=True, size_pt=14,
                 color=RGBColor(0x1E, 0x40, 0xAF),
                 alineacion=WD_ALIGN_PARAGRAPH.CENTER)
    _add_parrafo(doc, "", sangria=False)

    cuerpo = secciones.get("cuerpo", "").strip()
    for parrafo in cuerpo.split("\n"):
        linea = parrafo.strip()
        if not linea:
            _add_parrafo(doc, "", sangria=False); continue
        m = re.match(r"^\[H([123])\]\s*(.+)$", linea)
        if m:
            nivel = int(m.group(1)); titulo_h = m.group(2).strip()
            if nivel == 1:
                p = _add_parrafo(doc, titulo_h, sangria=False, bold=True,
                                 size_pt=12,
                                 alineacion=WD_ALIGN_PARAGRAPH.CENTER)
            elif nivel == 2:
                p = _add_parrafo(doc, titulo_h, sangria=False, bold=True,
                                 size_pt=12,
                                 alineacion=WD_ALIGN_PARAGRAPH.LEFT)
            else:
                p = _add_parrafo(doc, titulo_h, sangria=False, bold=True,
                                 size_pt=12,
                                 alineacion=WD_ALIGN_PARAGRAPH.LEFT)
                for r in p.runs:
                    r.italic = True
            # Mantiene el título junto con el siguiente párrafo
            p.paragraph_format.keep_with_next = True
            p.paragraph_format.keep_together = True
            continue
        _add_parrafo(doc, linea, sangria=True)

    refs = secciones.get("referencias", "").strip()
    if refs:
        doc.add_page_break()
        _add_parrafo(doc, "Referencias", sangria=False, bold=True, size_pt=12,
                     color=RGBColor(0x1E, 0x40, 0xAF),
                     alineacion=WD_ALIGN_PARAGRAPH.CENTER)
        for l in refs.split("\n"):
            l = l.strip()
            if l:
                _add_parrafo(doc, l, sangria=False, francesa=True)

    notas = secciones.get("notas", "").strip()
    if notas:
        doc.add_page_break()
        _add_parrafo(doc, "Notas de cambios", sangria=False, bold=True,
                     size_pt=12, color=RGBColor(0x1E, 0x40, 0xAF),
                     alineacion=WD_ALIGN_PARAGRAPH.CENTER)
        for l in notas.split("\n"):
            if l.strip():
                _add_parrafo(doc, l.strip(), sangria=False)

    buf = io.BytesIO(); doc.save(buf); return buf.getvalue()


def leer_docx(ruta: str) -> str:
    doc = Document(ruta)
    partes = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    for tabla in doc.tables:
        for fila in tabla.rows:
            for celda in fila.cells:
                t = (celda.text or "").strip()
                if t:
                    partes.append(t)
    return "\n\n".join(partes)


# ---------------------------- Almacenamiento ---------------------------------


def directorio_descargas() -> str:
    if platform == "android":
        try:
            from android.storage import primary_external_storage_path
            return os.path.join(primary_external_storage_path(), "Download")
        except Exception:
            return os.path.expanduser("~")
    return str(Path.home() / "Downloads")


def pedir_permisos_android():
    if platform != "android":
        return
    try:
        from android.permissions import Permission, request_permissions
        request_permissions([
            Permission.INTERNET,
            Permission.READ_EXTERNAL_STORAGE,
            Permission.WRITE_EXTERNAL_STORAGE,
        ])
    except Exception:
        pass


def detectar_modo_oscuro_sistema() -> str:
    """Devuelve 'Dark' o 'Light' según el sistema (en Android), si se puede."""
    if platform == "android":
        try:
            from jnius import autoclass
            Configuration = autoclass("android.content.res.Configuration")
            PythonActivity = autoclass(
                "org.kivy.android.PythonActivity"
            )
            night_mask = (
                PythonActivity.mActivity.getResources()
                .getConfiguration().uiMode
                & Configuration.UI_MODE_NIGHT_MASK
            )
            if night_mask == Configuration.UI_MODE_NIGHT_YES:
                return "Dark"
        except Exception:
            pass
    return "Light"


# ---------------------------- UI ---------------------------------------------


PRIMARIO = (30 / 255, 64 / 255, 175 / 255, 1)


class Tab(MDFloatLayout, MDTabsBase):
    """Contenedor para una pestaña en MDTabs."""


class HomeScreen(MDScreen):
    pass


class ResultsScreen(MDScreen):
    pass


class CorrectorApp(MDApp):
    title = "Corrector APA 7"

    def build(self):
        _marca_arranque("3_build_iniciado")
        self.theme_cls.primary_palette = "Blue"
        self.theme_cls.accent_palette = "Indigo"
        try:
            self.theme_cls.theme_style = detectar_modo_oscuro_sistema()
        except Exception as e:
            _marca_arranque("3a_tema_fallo", str(e))
            self.theme_cls.theme_style = "Light"
        # NOTA: M3 rompe MDSwitch en KivyMD 1.2.0 (KeyError 'thumb').
        # Usamos M2 que es estable y visualmente casi igual.
        self.theme_cls.material_style = "M2"
        _marca_arranque("3b_tema_ok")

        self.config_path = os.path.join(self.user_data_dir, "config.json")
        self._cargar_config()
        _marca_arranque("3c_config_ok")

        Window.softinput_mode = "below_target"
        try:
            pedir_permisos_android()
        except Exception as e:
            _marca_arranque("3d_permisos_fallo", str(e))
        _marca_arranque("3e_permisos_ok")

        self.sm = MDScreenManager()
        self.sm.add_widget(self._build_home())
        _marca_arranque("3f_home_ok")
        self.sm.add_widget(self._build_results())
        _marca_arranque("4_build_completo")
        return self.sm

    # ---- Configuración persistente (API key, modelo, tema) -----------------

    def _cargar_config(self):
        self.cfg = {
            "api_key": os.environ.get("GEMINI_API_KEY", ""),
            "modelo": GEMINI_DEFAULT,
        }
        try:
            if os.path.exists(self.config_path):
                with open(self.config_path, "r", encoding="utf-8") as f:
                    self.cfg.update(json.load(f))
        except Exception:
            pass

    def _guardar_config(self):
        try:
            os.makedirs(os.path.dirname(self.config_path), exist_ok=True)
            with open(self.config_path, "w", encoding="utf-8") as f:
                json.dump(self.cfg, f)
        except Exception:
            pass

    # ---- Home --------------------------------------------------------------

    def _build_home(self) -> HomeScreen:
        screen = HomeScreen(name="home")
        root = MDBoxLayout(orientation="vertical")

        toolbar = MDTopAppBar(
            title="Corrector APA 7",
            elevation=2,
            md_bg_color=PRIMARIO,
            specific_text_color=(1, 1, 1, 1),
            right_action_items=[
                ["theme-light-dark", lambda x: self._toggle_tema()],
                ["cog-outline", lambda x: self._abrir_ajustes()],
            ],
        )
        root.add_widget(toolbar)

        scroll = MDScrollView()
        body = MDBoxLayout(
            orientation="vertical", padding=dp(20), spacing=dp(16),
            adaptive_height=True, size_hint_y=None,
        )
        body.bind(minimum_height=body.setter("height"))

        body.add_widget(MDLabel(
            text="Pega tu texto académico o sube un .docx y lo corregimos al "
                 "formato APA 7.",
            theme_text_color="Secondary",
            size_hint_y=None, height=dp(60),
        ))

        # Botón grande: Subir documento
        body.add_widget(self._tarjeta_accion(
            icono="file-upload",
            titulo="Subir documento Word",
            subtitulo="Lee el .docx y carga el texto automáticamente",
            on_press=self._elegir_docx,
        ))

        # Tipo de documento
        self.tipo_btn = MDRaisedButton(
            text="Tipo: Ensayo / Trabajo académico",
            md_bg_color=PRIMARIO, size_hint_x=1,
            on_release=lambda *_: self._abrir_menu_tipo(),
        )
        body.add_widget(self.tipo_btn)
        self.tipo_doc = "ensayo o trabajo académico"

        # Caja de texto
        self.input_text = MDTextField(
            hint_text="Pega aquí tu texto…",
            multiline=True,
            mode="rectangle",
            size_hint_y=None, height=dp(220),
        )
        body.add_widget(self.input_text)

        # Switch notas
        notas_box = MDBoxLayout(
            orientation="horizontal", size_hint_y=None, height=dp(48),
            spacing=dp(12),
        )
        self.switch_notas = MDSwitch(active=True)
        notas_box.add_widget(self.switch_notas)
        notas_box.add_widget(MDLabel(
            text="Incluir notas de cambios en el Word",
            theme_text_color="Primary",
        ))
        body.add_widget(notas_box)

        # Botón corregir
        self.btn_corregir = MDRaisedButton(
            text="Corregir formato APA 7",
            icon="auto-fix",
            md_bg_color=PRIMARIO, size_hint_x=1,
            font_size=dp(18),
            on_release=lambda *_: self._iniciar_correccion(),
        )
        body.add_widget(self.btn_corregir)

        # Barra de progreso
        self.progress = MDProgressBar(
            value=0, size_hint_y=None, height=dp(6), opacity=0,
        )
        body.add_widget(self.progress)
        self.estado = MDLabel(
            text="", theme_text_color="Secondary",
            size_hint_y=None, height=dp(24), halign="center",
        )
        body.add_widget(self.estado)

        scroll.add_widget(body)
        root.add_widget(scroll)
        screen.add_widget(root)
        return screen

    def _tarjeta_accion(self, icono, titulo, subtitulo, on_press):
        card = MDCard(
            orientation="horizontal", padding=dp(16), spacing=dp(12),
            size_hint_y=None, height=dp(96), ripple_behavior=True,
            elevation=2, radius=[dp(14)] * 4,
        )
        card.bind(on_release=lambda *_: on_press())
        card.add_widget(MDIconButton(
            icon=icono, theme_text_color="Custom",
            text_color=PRIMARIO, icon_size=dp(36), disabled=True,
        ))
        col = MDBoxLayout(orientation="vertical")
        col.add_widget(MDLabel(text=titulo, font_style="H6", bold=True))
        col.add_widget(MDLabel(
            text=subtitulo, theme_text_color="Secondary",
        ))
        card.add_widget(col)
        return card

    def _abrir_menu_tipo(self):
        opciones = [
            "ensayo o trabajo académico",
            "artículo de investigación",
            "tesis o tesina",
            "informe / reporte",
            "monografía",
        ]
        items = [
            {
                "text": op.capitalize(),
                "viewclass": "OneLineListItem",
                "on_release": lambda x=op: self._set_tipo(x),
            }
            for op in opciones
        ]
        self._menu_tipo = MDDropdownMenu(
            caller=self.tipo_btn, items=items, width_mult=4,
        )
        self._menu_tipo.open()

    def _set_tipo(self, t):
        self.tipo_doc = t
        self.tipo_btn.text = f"Tipo: {t.capitalize()}"
        if hasattr(self, "_menu_tipo"):
            self._menu_tipo.dismiss()

    def _toggle_tema(self):
        self.theme_cls.theme_style = (
            "Dark" if self.theme_cls.theme_style == "Light" else "Light"
        )

    def _abrir_ajustes(self):
        contenido = MDBoxLayout(
            orientation="vertical", spacing=dp(12),
            size_hint_y=None, height=dp(140), padding=dp(8),
        )
        api_field = MDTextField(
            hint_text="API key de Gemini",
            text=self.cfg.get("api_key", ""),
            password=True,
        )
        modelo_field = MDTextField(
            hint_text="Modelo (gemini-2.0-flash, gemini-1.5-pro, …)",
            text=self.cfg.get("modelo", GEMINI_DEFAULT),
        )
        contenido.add_widget(api_field)
        contenido.add_widget(modelo_field)

        def guardar(*_):
            self.cfg["api_key"] = api_field.text.strip()
            self.cfg["modelo"] = modelo_field.text.strip() or GEMINI_DEFAULT
            self._guardar_config()
            dialog.dismiss()
            self._snack("Ajustes guardados.")

        dialog = MDDialog(
            title="Ajustes",
            type="custom",
            content_cls=contenido,
            buttons=[
                MDFlatButton(text="Cancelar",
                             on_release=lambda *_: dialog.dismiss()),
                MDFlatButton(text="Guardar",
                             theme_text_color="Custom",
                             text_color=PRIMARIO,
                             on_release=guardar),
            ],
        )
        dialog.open()

    # ---- File picker -------------------------------------------------------

    def _elegir_docx(self):
        try:
            from plyer import filechooser
        except Exception:
            self._snack("No se puede abrir el explorador de archivos.")
            return
        try:
            filechooser.open_file(
                title="Selecciona un archivo .docx",
                filters=[("Word", "*.docx")],
                on_selection=self._on_archivo_seleccionado,
            )
        except Exception as e:
            self._snack(f"No se pudo abrir el explorador: {e}")

    @mainthread
    def _on_archivo_seleccionado(self, seleccion):
        if not seleccion:
            return
        ruta = seleccion[0]
        try:
            texto = leer_docx(ruta)
            if not texto.strip():
                self._snack("El documento está vacío.")
                return
            self.input_text.text = texto
            self._snack(
                f"✅ Texto extraído ({len(texto.split())} palabras)."
            )
        except Exception as e:
            self._snack(f"No se pudo leer el .docx: {e}")

    # ---- Procesamiento -----------------------------------------------------

    def _iniciar_correccion(self):
        texto = (self.input_text.text or "").strip()
        if not texto:
            self._snack("Pega texto o sube un .docx primero.")
            return
        if not self.cfg.get("api_key"):
            self._snack("Configura tu API key en ajustes.")
            return

        self.btn_corregir.disabled = True
        self.progress.opacity = 1
        self.progress.value = 10
        self.estado.text = "Conectando con Gemini…"
        self._anim_progress()

        self._texto_original = texto

        hilo = threading.Thread(
            target=self._worker_correccion, args=(texto,), daemon=True,
        )
        hilo.start()

    def _anim_progress(self):
        def step(_dt):
            if self.progress.opacity == 0:
                return False
            v = self.progress.value
            if v < 90:
                self.progress.value = min(90, v + 4)
                return True
            return True
        Clock.schedule_interval(step, 0.4)

    def _worker_correccion(self, texto):
        try:
            salida = llamar_gemini_rest(
                api_key=self.cfg["api_key"],
                texto=texto,
                tipo_doc=self.tipo_doc,
                modelo=self.cfg.get("modelo", GEMINI_DEFAULT),
                con_notas=bool(self.switch_notas.active),
            )
            secciones = parsear_secciones(salida)
            self._on_correccion_ok(secciones)
        except Exception as e:
            self._on_correccion_error(str(e))

    @mainthread
    def _on_correccion_ok(self, secciones):
        self.progress.value = 100
        self.estado.text = "✅ Corrección completada."
        self.btn_corregir.disabled = False
        self.progress.opacity = 0
        self._secciones = secciones
        self._mostrar_resultados()

    @mainthread
    def _on_correccion_error(self, mensaje):
        self.progress.opacity = 0
        self.btn_corregir.disabled = False
        self.estado.text = ""
        self._snack(mensaje)

    # ---- Resultados --------------------------------------------------------

    def _build_results(self) -> ResultsScreen:
        screen = ResultsScreen(name="results")
        root = MDBoxLayout(orientation="vertical")

        toolbar = MDTopAppBar(
            title="Resultado APA 7",
            elevation=2,
            md_bg_color=PRIMARIO,
            specific_text_color=(1, 1, 1, 1),
            left_action_items=[
                ["arrow-left", lambda x: self._volver_home()],
            ],
            right_action_items=[
                ["download", lambda x: self._descargar_word()],
                ["theme-light-dark", lambda x: self._toggle_tema()],
            ],
        )
        root.add_widget(toolbar)

        self.tabs = MDTabs(
            background_color=PRIMARIO,
            text_color_active=(1, 1, 1, 1),
            text_color_normal=(1, 1, 1, 0.6),
            indicator_color=(1, 1, 1, 1),
        )
        self.tab_cuerpo = Tab(title="Cuerpo")
        self.tab_refs = Tab(title="Referencias")
        self.tab_stats = Tab(title="Estadísticas")
        self.tab_tips = Tab(title="Tips")
        for t in (self.tab_cuerpo, self.tab_refs, self.tab_stats, self.tab_tips):
            self.tabs.add_widget(t)
        root.add_widget(self.tabs)
        screen.add_widget(root)
        return screen

    def _mostrar_resultados(self):
        for tab in (self.tab_cuerpo, self.tab_refs, self.tab_stats,
                    self.tab_tips):
            tab.clear_widgets()

        s = self._secciones

        # CUERPO
        self.tab_cuerpo.add_widget(
            self._scroll_texto(self._formatear_cuerpo_markup(
                s.get("titulo", ""), s.get("cuerpo", "")
            ), markup=True)
        )

        # REFERENCIAS
        refs = s.get("referencias", "").strip()
        if refs:
            txt = "\n\n".join(f"• {l.strip()}" for l in refs.split("\n")
                              if l.strip())
        else:
            txt = "No se generaron referencias."
        self.tab_refs.add_widget(self._scroll_texto(txt))

        # ESTADÍSTICAS
        self.tab_stats.add_widget(self._build_stats_view(s))

        # TIPS
        self.tab_tips.add_widget(self._build_tips_view(s))

        self.sm.current = "results"

    def _formatear_cuerpo_markup(self, titulo, cuerpo):
        partes = []
        if titulo:
            partes.append(
                f"[size=22sp][b][color=1E40AF]{titulo}[/color][/b][/size]\n"
            )
        for linea in (cuerpo or "").split("\n"):
            l = linea.strip()
            if not l:
                partes.append(""); continue
            m = re.match(r"^\[H([123])\]\s*(.+)$", l)
            if m:
                nivel = int(m.group(1)); txt = m.group(2)
                if nivel == 1:
                    partes.append(f"\n[size=20sp][b]{txt}[/b][/size]\n")
                elif nivel == 2:
                    partes.append(f"\n[size=18sp][b]{txt}[/b][/size]\n")
                else:
                    partes.append(f"\n[size=16sp][b][i]{txt}[/i][/b][/size]\n")
            else:
                partes.append(l)
        return "\n".join(partes)

    def _scroll_texto(self, texto, markup=False):
        scroll = MDScrollView()
        lbl = MDLabel(
            text=texto or "—",
            markup=markup,
            size_hint_y=None,
            padding=(dp(16), dp(16)),
            halign="left",
        )
        lbl.bind(
            width=lambda *_: setattr(lbl, "text_size",
                                     (lbl.width - dp(32), None)),
            texture_size=lambda *_: setattr(lbl, "height",
                                            lbl.texture_size[1] + dp(32)),
        )
        scroll.add_widget(lbl)
        return scroll

    def _build_stats_view(self, secciones):
        scroll = MDScrollView()
        box = MDBoxLayout(
            orientation="vertical", padding=dp(16), spacing=dp(10),
            adaptive_height=True, size_hint_y=None,
        )
        box.bind(minimum_height=box.setter("height"))

        original = getattr(self, "_texto_original", "")
        cuerpo = secciones.get("cuerpo", "")
        refs = secciones.get("referencias", "")
        cambios = contar_cambios_palabras(original, cuerpo)
        citas_o = contar_citas(original); citas_c = contar_citas(cuerpo)
        n_refs = contar_referencias(refs)

        box.add_widget(MDLabel(
            text="[b]Métricas de la corrección[/b]",
            markup=True, font_style="H6",
            size_hint_y=None, height=dp(28),
        ))
        for label, valor in [
            ("Citas en el texto", f"{citas_c} (Δ {citas_c - citas_o:+d})"),
            ("Referencias generadas", str(n_refs)),
            ("Palabras corregidas", str(cambios["total"])),
            ("Modificadas / añadidas / eliminadas",
             f"{cambios['modificadas']} / {cambios['anadidas']} / "
             f"{cambios['eliminadas']}"),
            ("Palabras totales",
             f"{cambios['palabras_corregido']} "
             f"({cambios['palabras_corregido'] - cambios['palabras_original']:+d})"),
        ]:
            box.add_widget(self._fila_metrica(label, valor))

        # Distribución por tipo de fuente
        stats = estadisticas_referencias(refs)
        box.add_widget(MDLabel(
            text="[b]Distribución de fuentes[/b]",
            markup=True, font_style="H6",
            size_hint_y=None, height=dp(34),
        ))
        if stats["total"] == 0:
            box.add_widget(MDLabel(
                text="No hay referencias para analizar.",
                theme_text_color="Secondary",
                size_hint_y=None, height=dp(24),
            ))
        else:
            max_v = max(stats["por_tipo"].values())
            for tipo, n in sorted(stats["por_tipo"].items(),
                                  key=lambda x: -x[1]):
                box.add_widget(self._barra_tipo(
                    tipo, n, max_v, stats["total"]))

        scroll.add_widget(box)
        return scroll

    def _fila_metrica(self, etiqueta, valor):
        row = MDBoxLayout(
            orientation="horizontal", size_hint_y=None, height=dp(32),
        )
        row.add_widget(MDLabel(text=etiqueta, theme_text_color="Secondary"))
        row.add_widget(MDLabel(
            text=valor, halign="right", bold=True,
        ))
        return row

    def _barra_tipo(self, tipo, n, max_v, total):
        cont = MDBoxLayout(
            orientation="vertical", size_hint_y=None, height=dp(48),
            spacing=dp(2),
        )
        head = MDBoxLayout(
            orientation="horizontal", size_hint_y=None, height=dp(20),
        )
        head.add_widget(MDLabel(text=tipo))
        head.add_widget(MDLabel(
            text=f"{n}  ({round(100 * n / total)}%)", halign="right",
        ))
        cont.add_widget(head)
        bar = MDProgressBar(
            value=int(100 * n / max(max_v, 1)),
            size_hint_y=None, height=dp(8),
            color=PRIMARIO,
        )
        cont.add_widget(bar)
        return cont

    def _build_tips_view(self, secciones):
        scroll = MDScrollView()
        box = MDBoxLayout(
            orientation="vertical", padding=dp(16), spacing=dp(8),
            adaptive_height=True, size_hint_y=None,
        )
        box.bind(minimum_height=box.setter("height"))

        a = analizar_legibilidad(secciones.get("cuerpo", "")
                                 or getattr(self, "_texto_original", ""))
        box.add_widget(MDLabel(
            text="[b]Análisis de redacción[/b]",
            markup=True, font_style="H6",
            size_hint_y=None, height=dp(34),
        ))
        for label, valor in [
            ("Legibilidad (Fernández-Huerta)",
             f"{a['fh']:.0f}/100  ·  {a['nivel']}"),
            ("Promedio de palabras por oración", f"{a['F']:.1f}"),
            ("Diversidad léxica", f"{a['diversidad']:.0f}%"),
            ("Oraciones largas (>30 palabras)", str(a["oraciones_largas"])),
        ]:
            box.add_widget(self._fila_metrica(label, valor))

        box.add_widget(MDLabel(
            text="[b]Sugerencias[/b]",
            markup=True, font_style="H6",
            size_hint_y=None, height=dp(34),
        ))
        for c in a["consejos"]:
            box.add_widget(MDLabel(
                text=f"• {c}",
                size_hint_y=None,
            ))
            # auto-altura por etiqueta
            box.children[0].bind(
                width=lambda lbl, w: setattr(
                    lbl, "text_size", (w - dp(32), None)),
                texture_size=lambda lbl, ts: setattr(
                    lbl, "height", ts[1] + dp(8)),
            )
        scroll.add_widget(box)
        return scroll

    def _volver_home(self):
        self.sm.current = "home"

    # ---- Descarga Word -----------------------------------------------------

    def _descargar_word(self):
        try:
            data = construir_docx(self._secciones)
            destino = directorio_descargas()
            os.makedirs(destino, exist_ok=True)
            base = (self._secciones.get("titulo") or "documento_apa")
            nombre = re.sub(r"[^\w\-]+", "_", base.strip())[:60] or "documento_apa"
            ruta = os.path.join(destino, f"{nombre}.docx")
            i = 1
            while os.path.exists(ruta):
                ruta = os.path.join(destino, f"{nombre}_{i}.docx")
                i += 1
            with open(ruta, "wb") as f:
                f.write(data)
            self._snack(f"📥 Guardado en: {ruta}")
        except Exception as e:
            self._snack(f"No se pudo guardar el .docx: {e}")

    # ---- Util --------------------------------------------------------------

    def _snack(self, mensaje):
        try:
            Snackbar(text=mensaje, duration=3).open()
        except Exception:
            print(mensaje)


if __name__ == "__main__":
    try:
        CorrectorApp().run()
    except Exception:
        _escribir_crash(*sys.exc_info())
        raise
