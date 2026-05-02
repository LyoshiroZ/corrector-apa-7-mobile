"""Plantillas APA 7 prediseñadas para descarga directa en Word."""

from __future__ import annotations

import io

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


PLANTILLAS_APA: dict = {
    "Ensayo": {
        "descripcion": "Estructura clásica para un ensayo académico breve.",
        "secciones": [
            ("H1", "Introducción",
             "Presenta el tema, justifica su importancia y enuncia la tesis "
             "que defenderás a lo largo del ensayo. Cierra con un breve "
             "anticipo de los argumentos que desarrollarás."),
            ("H1", "Desarrollo", ""),
            ("H2", "Argumento 1",
             "Expón tu primer argumento con claridad. Apóyalo con citas en "
             "formato APA 7, por ejemplo (Apellido, año)."),
            ("H2", "Argumento 2",
             "Continúa con el segundo argumento, conectándolo con el "
             "anterior. Usa evidencia y ejemplos concretos."),
            ("H2", "Contraargumento",
             "Reconoce posibles objeciones y responde a ellas para reforzar "
             "tu postura."),
            ("H1", "Conclusiones",
             "Sintetiza los argumentos principales, reformula la tesis y "
             "propone una reflexión final o una pregunta abierta."),
        ],
    },
    "Artículo de investigación": {
        "descripcion": "Formato IMRyD para un artículo académico.",
        "secciones": [
            ("H1", "Resumen",
             "[Resumen en un solo párrafo de 150-250 palabras. Describe el "
             "problema, el método, los resultados principales y la "
             "conclusión.]"),
            ("H1", "Palabras clave",
             "[3 a 5 palabras clave separadas por punto y coma.]"),
            ("H1", "Introducción",
             "Contextualiza el problema, revisa brevemente la literatura "
             "relevante y formula los objetivos o hipótesis."),
            ("H1", "Método", ""),
            ("H2", "Participantes",
             "Describe la muestra: número, características demográficas y "
             "criterios de inclusión."),
            ("H2", "Instrumentos",
             "Detalla los instrumentos, escalas o materiales empleados."),
            ("H2", "Procedimiento",
             "Explica paso a paso cómo se recogieron los datos."),
            ("H2", "Análisis de datos",
             "Indica las técnicas estadísticas o cualitativas utilizadas."),
            ("H1", "Resultados",
             "Presenta los hallazgos sin interpretarlos. Apóyate en tablas "
             "y figuras numeradas según APA 7."),
            ("H1", "Discusión",
             "Interpreta los resultados a la luz de la literatura previa, "
             "señala limitaciones y propone líneas futuras."),
            ("H1", "Conclusiones",
             "Sintetiza los aportes principales del estudio."),
        ],
    },
    "Informe de investigación": {
        "descripcion": "Estructura completa para informes académicos.",
        "secciones": [
            ("H1", "Resumen ejecutivo",
             "Síntesis breve del problema, metodología, hallazgos y "
             "recomendaciones."),
            ("H1", "Introducción",
             "Contexto general del informe y su propósito."),
            ("H2", "Antecedentes",
             "Describe el origen del problema y los estudios previos."),
            ("H2", "Justificación",
             "Explica la pertinencia y utilidad del informe."),
            ("H1", "Objetivos", ""),
            ("H2", "Objetivo general",
             "[Redacta el objetivo general en una sola oración con verbo "
             "en infinitivo.]"),
            ("H2", "Objetivos específicos",
             "[Lista 3 a 5 objetivos específicos.]"),
            ("H1", "Metodología",
             "Tipo de estudio, técnicas e instrumentos utilizados."),
            ("H1", "Resultados",
             "Presentación organizada de los hallazgos."),
            ("H1", "Discusión",
             "Interpretación de los resultados frente al marco teórico."),
            ("H1", "Conclusiones y recomendaciones",
             "Conclusiones derivadas y recomendaciones prácticas."),
        ],
    },
    "Tesis / Monografía": {
        "descripcion": "Estructura por capítulos con páginas preliminares.",
        "secciones": [
            ("H1", "Dedicatoria",
             "[Texto breve y personal dirigido a quien dedicas el trabajo.]"),
            ("H1", "Agradecimientos",
             "[Personas e instituciones que apoyaron el proceso.]"),
            ("H1", "Resumen",
             "[Resumen general de la investigación, 200-300 palabras.]"),
            ("H1", "Capítulo I: Planteamiento del problema", ""),
            ("H2", "Descripción de la realidad problemática",
             "Describe el contexto y el problema."),
            ("H2", "Formulación del problema",
             "Pregunta general y preguntas específicas de investigación."),
            ("H2", "Objetivos",
             "Objetivo general y objetivos específicos."),
            ("H2", "Justificación",
             "Importancia teórica, social y práctica del estudio."),
            ("H1", "Capítulo II: Marco teórico", ""),
            ("H2", "Antecedentes de la investigación",
             "Estudios nacionales e internacionales relacionados."),
            ("H2", "Bases teóricas",
             "Teorías y conceptos que sustentan la investigación."),
            ("H2", "Definición de términos básicos",
             "Glosario de términos clave."),
            ("H1", "Capítulo III: Metodología", ""),
            ("H2", "Tipo y diseño de investigación",
             "Enfoque, tipo, nivel y diseño del estudio."),
            ("H2", "Población y muestra",
             "Universo, muestra y criterios de selección."),
            ("H2", "Técnicas e instrumentos",
             "Herramientas de recolección de datos."),
            ("H1", "Capítulo IV: Resultados",
             "Presentación e interpretación de los datos obtenidos."),
            ("H1", "Capítulo V: Conclusiones y recomendaciones",
             "Conclusiones del estudio y recomendaciones para futuros "
             "trabajos."),
        ],
    },
    "Reseña crítica": {
        "descripcion": "Para reseñar un libro, artículo o capítulo.",
        "secciones": [
            ("H1", "Datos bibliográficos",
             "[Cita en APA 7 de la obra reseñada.]"),
            ("H1", "Presentación de la obra",
             "Información sobre el autor, el contexto y el propósito de la "
             "obra."),
            ("H1", "Resumen del contenido",
             "Síntesis objetiva de las ideas principales del texto."),
            ("H1", "Análisis crítico",
             "Valoración fundamentada: fortalezas, debilidades y aportes a "
             "su campo."),
            ("H1", "Conclusión",
             "Recomendación final y reflexión personal."),
        ],
    },
    "Proyecto de investigación": {
        "descripcion": "Anteproyecto previo a la tesis o investigación.",
        "secciones": [
            ("H1", "Título del proyecto",
             "[Título tentativo de la investigación.]"),
            ("H1", "Planteamiento del problema",
             "Descripción del problema y su contexto."),
            ("H1", "Pregunta de investigación",
             "[Pregunta principal y subpreguntas si las hay.]"),
            ("H1", "Objetivos", ""),
            ("H2", "Objetivo general",
             "[Verbo en infinitivo + qué + para qué.]"),
            ("H2", "Objetivos específicos", "[Tres a cinco objetivos.]"),
            ("H1", "Justificación",
             "Importancia, viabilidad y aporte del estudio."),
            ("H1", "Marco teórico preliminar",
             "Principales teorías y antecedentes que orientarán la "
             "investigación."),
            ("H1", "Metodología propuesta",
             "Enfoque, diseño, técnicas e instrumentos previstos."),
            ("H1", "Cronograma",
             "[Tabla o lista con las actividades y fechas estimadas.]"),
            ("H1", "Referencias preliminares",
             "Listado inicial de fuentes consultadas."),
        ],
    },
}


def _aplicar_fuente(run, *, size_pt: float = 12, bold: bool = False, color=None):
    run.font.name = "Times New Roman"
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = rPr.makeelement(qn("w:rFonts"), {})
        rPr.append(rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rFonts.set(qn(attr), "Times New Roman")
    run.font.size = Pt(size_pt)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def _add_parrafo(doc, texto: str, *, sangria: bool = True, francesa: bool = False,
                 size_pt: float = 12, bold: bool = False, color=None,
                 alineacion=None):
    p = doc.add_paragraph()
    if alineacion is not None:
        p.alignment = alineacion
    pf = p.paragraph_format
    if francesa:
        pf.left_indent = Cm(1.27)
        pf.first_line_indent = Cm(-1.27)
    elif sangria:
        pf.first_line_indent = Cm(1.27)
    pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)
    run = p.add_run(texto)
    _aplicar_fuente(run, size_pt=size_pt, bold=bold, color=color)
    return p


def construir_plantilla_apa(tipo: str) -> bytes:
    """Genera un .docx con carátula APA 7 y estructura preconstruida."""
    plantilla = PLANTILLAS_APA[tipo]
    doc = Document()

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    rPr = normal.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = rPr.makeelement(qn("w:rFonts"), {})
        rPr.append(rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rFonts.set(qn(attr), "Times New Roman")

    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    # ---- Carátula APA ----
    for _ in range(6):
        _add_parrafo(doc, "", sangria=False)

    _add_parrafo(
        doc, f"[Título del {tipo.lower()}]",
        sangria=False, bold=True, size_pt=14,
        color=RGBColor(0x2D, 0x50, 0x16),
        alineacion=WD_ALIGN_PARAGRAPH.CENTER,
    )
    _add_parrafo(doc, "", sangria=False)
    for linea in [
        "[Nombre completo del estudiante]",
        "[Facultad / Escuela]",
        "[Universidad]",
        "[Asignatura: nombre del curso]",
        "[Docente: Apellido, Inicial.]",
        "[Fecha: día de mes de año]",
    ]:
        _add_parrafo(doc, linea, sangria=False,
                     alineacion=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_page_break()

    # ---- Cuerpo según secciones ----
    for nivel, titulo, cuerpo in plantilla["secciones"]:
        if nivel == "H1":
            p = _add_parrafo(
                doc, titulo, sangria=False, bold=True, size_pt=12,
                alineacion=WD_ALIGN_PARAGRAPH.CENTER,
            )
        elif nivel == "H2":
            p = _add_parrafo(
                doc, titulo, sangria=False, bold=True, size_pt=12,
                alineacion=WD_ALIGN_PARAGRAPH.LEFT,
            )
        else:
            p = _add_parrafo(
                doc, titulo, sangria=False, bold=True, size_pt=12,
                alineacion=WD_ALIGN_PARAGRAPH.LEFT,
            )
            for r in p.runs:
                r.italic = True
        p.paragraph_format.keep_with_next = True
        p.paragraph_format.keep_together = True

        if cuerpo:
            _add_parrafo(doc, cuerpo, sangria=True)
        _add_parrafo(doc, "", sangria=False)

    # ---- Página de Referencias ----
    doc.add_page_break()
    _add_parrafo(
        doc, "Referencias", sangria=False, bold=True, size_pt=12,
        color=RGBColor(0x2D, 0x50, 0x16),
        alineacion=WD_ALIGN_PARAGRAPH.CENTER,
    )
    _add_parrafo(
        doc,
        "[Apellido, A. A. (año). Título del trabajo. Editorial. "
        "https://doi.org/xxxx]",
        sangria=False, francesa=True,
    )
    _add_parrafo(
        doc,
        "[Apellido, B. B., y Apellido, C. C. (año). Título del artículo. "
        "Nombre de la Revista, vol(núm), pp-pp.]",
        sangria=False, francesa=True,
    )

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
